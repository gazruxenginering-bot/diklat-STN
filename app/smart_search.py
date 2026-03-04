"""
Smart Document Search Module dengan Chroma Vector Database
Mengambil dokumen dari Google Drive, parse, store di Chroma, dan search dengan semantic similarity
"""

import os
import json
import re
from typing import List, Dict, Tuple
from datetime import datetime
import numpy as np
from io import BytesIO

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

from google.auth.transport.requests import Request
from google.oauth2.service_account import Credentials
from googleapiclient import discovery, errors

# Import Chroma integration
try:
    from .chroma_integration import get_vector_store
    CHROMA_AVAILABLE = True
except ImportError:
    CHROMA_AVAILABLE = False
    print("⚠️  Chroma not available, will use fallback search")


class DocumentProcessor:
    """Process berbagai tipe dokumen (PDF, Word, Text)"""
    
    @staticmethod
    def extract_text_from_pdf(file_obj) -> str:
        """Extract text dari PDF"""
        if not PdfReader:
            return ""
        
        try:
            pdf_reader = PdfReader(file_obj)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
            return text
        except Exception as e:
            print(f"❌ Error extracting PDF: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_obj) -> str:
        """Extract text dari Word document"""
        if not Document:
            return ""
        
        try:
            doc = Document(file_obj)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return text
        except Exception as e:
            print(f"❌ Error extracting DOCX: {e}")
            return ""
    
    @staticmethod
    def extract_text_from_txt(file_obj) -> str:
        """Extract text dari text file"""
        try:
            content = file_obj.read()
            if isinstance(content, bytes):
                return content.decode('utf-8', errors='ignore')
            return content
        except Exception as e:
            print(f"❌ Error reading TXT: {e}")
            return ""


class GoogleDriveDocumentManager:
    """Manage documents dari Google Drive"""
    
    def __init__(self, credentials_path: str = None):
        self.credentials = None
        self.drive_service = None
        self.credentials_path = credentials_path or 'credentials.json'
        self._initialize_drive_service()
    
    def _initialize_drive_service(self):
        """Initialize Google Drive API service"""
        try:
            if os.path.exists(self.credentials_path):
                self.credentials = Credentials.from_service_account_file(
                    self.credentials_path,
                    scopes=['https://www.googleapis.com/auth/drive.readonly']
                )
                self.drive_service = discovery.build('drive', 'v3', credentials=self.credentials)
                print("✅ Google Drive Service initialized")
        except Exception as e:
            print(f"⚠️  Could not initialize Google Drive: {e}")
    
    def search_files(self, query: str, file_types: List[str] = None, limit: int = 10) -> List[Dict]:
        """
        Search files di Google Drive
        
        Args:
            query: Search query
            file_types: List of MIME types (e.g., ['application/pdf', 'application/vnd.openxmlformats'])
            limit: Max results
        
        Returns:
            List of file info
        """
        if not self.drive_service:
            return []
        
        try:
            # Build query
            q_parts = [f"name contains '{query}'", "trashed=false"]
            
            if file_types:
                mime_query = " or ".join([f"mimeType='{mt}'" for mt in file_types])
                q_parts.append(f"({mime_query})")
            
            query_str = " and ".join(q_parts)
            
            results = self.drive_service.files().list(
                q=query_str,
                spaces='drive',
                fields='files(id, name, mimeType, modifiedTime, webViewLink, parents)',
                pageSize=limit,
                orderBy='modifiedTime desc'
            ).execute()
            
            return results.get('files', [])
        except errors.HttpError as e:
            print(f"❌ Google Drive API error: {e}")
            return []
    
    def get_file_content(self, file_id: str) -> Tuple[str, str]:
        """
        Get file content dari Google Drive
        
        Returns:
            (content, mime_type)
        """
        if not self.drive_service:
            return "", ""
        
        try:
            # Get file metadata
            file_metadata = self.drive_service.files().get(
                fileId=file_id,
                fields='mimeType, name'
            ).execute()
            
            mime_type = file_metadata.get('mimeType', '')
            name = file_metadata.get('name', '')
            
            # Download file
            request = self.drive_service.files().get_media(fileId=file_id)
            file_obj = BytesIO(request.execute())
            
            # Extract text based on MIME type
            if 'pdf' in mime_type:
                text = DocumentProcessor.extract_text_from_pdf(file_obj)
            elif 'wordprocessingml' in mime_type or 'document' in mime_type:
                text = DocumentProcessor.extract_text_from_docx(file_obj)
            elif 'text' in mime_type:
                text = DocumentProcessor.extract_text_from_txt(file_obj)
            else:
                text = ""
            
            return text, mime_type
        except Exception as e:
            print(f"❌ Error getting file content: {e}")
            return "", ""


class TextChunker:
    """Split text menjadi chunks untuk processing"""
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        """
        Split text menjadi chunks dengan overlap
        
        Args:
            text: Full text to chunk
            chunk_size: Size of each chunk (characters)
            overlap: Overlap between chunks
        
        Returns:
            List of chunks
        """
        if not text:
            return []
        
        chunks = []
        step = chunk_size - overlap
        
        for i in range(0, len(text), step):
            chunk = text[i:i + chunk_size]
            if chunk.strip():
                chunks.append(chunk)
            
            if i + chunk_size >= len(text):
                break
        
        return chunks


class SimpleSemanticSearch:
    """Simple semantic search menggunakan keyword dan similarity"""
    
    @staticmethod
    def calculate_similarity(text1: str, text2: str) -> float:
        """Calculate similarity antara dua text menggunakan word overlap"""
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        if not words1 or not words2:
            return 0.0
        
        intersection = len(words1 & words2)
        union = len(words1 | words2)
        
        return intersection / union if union > 0 else 0.0
    
    @staticmethod
    def search_similar_chunks(query: str, chunks: List[str], top_k: int = 3) -> List[Tuple[str, float]]:
        """
        Search similar chunks to query
        
        Returns:
            List of (chunk, similarity_score) tuples
        """
        scores = []
        
        for chunk in chunks:
            score = SimpleSemanticSearch.calculate_similarity(query, chunk)
            scores.append((chunk, score))
        
        # Sort by score descending
        scores.sort(key=lambda x: x[1], reverse=True)
        
        return scores[:top_k]


class SmartDocumentSearch:
    """Main class untuk smart document search dengan RAG"""
    
    def __init__(self, credentials_path: str = None):
        self.drive_manager = GoogleDriveDocumentManager(credentials_path)
        self.document_cache = {}  # Cache dokumen yang sudah didownload
    
    def search_and_retrieve(self, 
                            query: str, 
                            search_limit: int = 5,
                            chunk_size: int = 1000,
                            top_chunks: int = 3) -> Dict:
        """
        Search dokumen relevan dan retrieve relevant chunks
        
        Args:
            query: User question/search term
            search_limit: Max docs to search
            chunk_size: Size of text chunks
            top_chunks: Top K chunks to return per document
        
        Returns:
            {
                'query': original query,
                'results': [
                    {
                        'file_name': str,
                        'file_id': str,
                        'mime_type': str,
                        'chunks': [
                            {'text': str, 'similarity': float}
                        ],
                        'web_view_link': str
                    }
                ]
            }
        """
        # Search files di Google Drive
        file_types = [
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'text/plain'
        ]
        
        found_files = self.drive_manager.search_files(query, file_types, search_limit)
        
        results = []
        
        for file_info in found_files:
            try:
                file_id = file_info.get('id')
                file_name = file_info.get('name')
                mime_type = file_info.get('mimeType', '')
                web_view_link = file_info.get('webViewLink', '')
                
                # Get file content (dengan caching)
                if file_id not in self.document_cache:
                    content, _ = self.drive_manager.get_file_content(file_id)
                    self.document_cache[file_id] = content
                else:
                    content = self.document_cache[file_id]
                
                if not content:
                    continue
                
                # Chunk text
                chunks = TextChunker.chunk_text(content, chunk_size)
                
                # Find relevant chunks
                relevant_chunks = SimpleSemanticSearch.search_similar_chunks(query, chunks, top_chunks)
                
                if relevant_chunks:
                    result = {
                        'file_name': file_name,
                        'file_id': file_id,
                        'mime_type': mime_type,
                        'web_view_link': web_view_link,
                        'chunks': [
                            {
                                'text': chunk,
                                'similarity': round(score, 3)
                            }
                            for chunk, score in relevant_chunks
                        ]
                    }
                    results.append(result)
            except Exception as e:
                print(f"❌ Error processing file {file_info.get('name')}: {e}")
                continue
        
        return {
            'query': query,
            'results': results,
            'total_files_found': len(found_files),
            'total_relevant_chunks': sum(len(r['chunks']) for r in results)
        }
    
    def format_context_for_ai(self, search_results: Dict) -> str:
        """Format search results untuk di-pass ke AI model"""
        context = "## Dokumen Relevan:\n\n"
        
        for i, result in enumerate(search_results.get('results', []), 1):
            context += f"### {i}. {result['file_name']}\n"
            context += f"Sumber: {result['web_view_link']}\n\n"
            
            for j, chunk in enumerate(result['chunks'], 1):
                context += f"**Bagian {j}** (Relevance: {chunk['similarity']}):\n"
                context += f"{chunk['text'][:500]}...\n\n" if len(chunk['text']) > 500 else f"{chunk['text']}\n\n"
        
        return context
    
    def clear_cache(self):
        """Clear document cache"""
        self.document_cache.clear()
        print("✅ Document cache cleared")


class ChromaDocumentSearch:
    """Document search menggunakan Chroma Vector Database untuk semantic similarity"""
    
    def __init__(self, credentials_path: str = None):
        """
        Initialize ChromaDocumentSearch
        
        Args:
            credentials_path: Path to Google Drive credentials.json
        """
        self.drive_manager = GoogleDriveDocumentManager(credentials_path)
        self.vector_store = get_vector_store() if CHROMA_AVAILABLE else None
        
        if not self.vector_store:
            print("❌ Chroma Vector Store not available")
    
    def index_document_from_drive(self, drive_file_id: str, drive_file_name: str) -> bool:
        """
        Download dokumen dari Google Drive dan index ke Chroma
        
        Args:
            drive_file_id: Google Drive file ID
            drive_file_name: File name
        
        Returns:
            Success status
        """
        if not self.vector_store:
            print("❌ Vector store not available")
            return False
        
        try:
            # Download content
            content, mime_type = self.drive_manager.get_file_content(drive_file_id)
            
            if not content:
                print(f"⚠️  Could not extract content from {drive_file_name}")
                return False
            
            # Chunk text
            chunks = TextChunker.chunk_text(content, chunk_size=1000, overlap=100)
            
            if not chunks:
                print(f"⚠️  No chunks generated for {drive_file_name}")
                return False
            
            # Add to Chroma
            metadata = {
                "mime_type": mime_type,
                "indexed_date": datetime.utcnow().isoformat(),
                "source": "google_drive"
            }
            
            success = self.vector_store.add_document_chunks(
                file_id=drive_file_id,
                file_name=drive_file_name,
                chunks=chunks,
                metadata=metadata
            )
            
            return success
        
        except Exception as e:
            print(f"❌ Error indexing document {drive_file_name}: {e}")
            return False
    
    def search(self, 
               query: str,
               search_limit: int = 5,
               results_limit: int = 10) -> Dict:
        """
        Semantic search menggunakan Chroma vector database
        
        Args:
            query: User question/search query
            search_limit: Max documents to return
            results_limit: Max chunks per document
        
        Returns:
            {
                'query': str,
                'results': [
                    {
                        'file_id': str,
                        'file_name': str,
                        'chunks': [
                            {
                                'text': str,
                                'similarity': float,
                                'chunk_index': int
                            }
                        ]
                    }
                ],
                'total_results': int
            }
        """
        if not self.vector_store:
            return {'query': query, 'results': [], 'total_results': 0}
        
        try:
            return self.vector_store.search_documents(query, search_limit, results_limit)
        except Exception as e:
            print(f"❌ Error searching documents: {e}")
            return {'query': query, 'results': [], 'total_results': 0}
    
    def update_document(self, drive_file_id: str, drive_file_name: str) -> bool:
        """
        Update dokumen yang sudah ada di Chroma
        
        Args:
            drive_file_id: Google Drive file ID
            drive_file_name: File name
        
        Returns:
            Success status
        """
        if not self.vector_store:
            return False
        
        try:
            # Delete old version
            self.vector_store.delete_document(drive_file_id)
            
            # Index new version
            return self.index_document_from_drive(drive_file_id, drive_file_name)
        
        except Exception as e:
            print(f"❌ Error updating document: {e}")
            return False
    
    def delete_document(self, drive_file_id: str) -> bool:
        """Delete dokumen dari Chroma"""
        if not self.vector_store:
            return False
        
        return self.vector_store.delete_document(drive_file_id)
    
    def get_stats(self) -> Dict:
        """Get Chroma collection statistics"""
        if not self.vector_store:
            return {}
        
        return self.vector_store.get_collection_stats()
    
    def format_context_for_ai(self, search_results: Dict) -> str:
        """
        Format Chroma search results untuk di-pass ke AI model
        
        Args:
            search_results: Results dari search()
        
        Returns:
            Formatted context string
        """
        context = "## Dokumen Relevan dari Knowledge Base:\n\n"
        
        for i, result in enumerate(search_results.get('results', []), 1):
            context += f"### {i}. {result['file_name']}\n"
            context += f"Document ID: {result['file_id']}\n\n"
            
            for j, chunk in enumerate(result['chunks'], 1):
                relevance_pct = int(chunk['similarity'] * 100)
                context += f"**Bagian {j}** (Relevansi: {relevance_pct}%):\n"
                
                # Truncate long chunks
                text = chunk['text']
                if len(text) > 800:
                    text = text[:800] + "..."
                
                context += f"{text}\n\n"
        
        return context

